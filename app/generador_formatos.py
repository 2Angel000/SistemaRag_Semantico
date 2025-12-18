from datetime import datetime
from docx import Document
import os


def reemplazar_en_documento(doc, mapping: dict[str, str]):
    """Reemplaza texto en párrafos y tablas, manejando casos donde Word parte los runs."""
    # Párrafos
    for paragraph in doc.paragraphs:
        for key, value in mapping.items():
            if key in paragraph.text:
                nuevo_texto = paragraph.text.replace(key, value)
                # Dejar un solo run con el texto nuevo
                while len(paragraph.runs) > 1:
                    r = paragraph.runs[1]
                    r._element.getparent().remove(r._element)
                paragraph.runs[0].text = nuevo_texto

    # Celdas de tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in mapping.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, value)


class GeneradorFormatosCalificaciones:
    def __init__(self, sistema_rag):
        self.sistema = sistema_rag
        # base_path = carpeta raíz del proyecto (un nivel arriba de app/)
        self.base_path = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

    def generar_formato_calificaciones(self, matricula_or_data):
        # 1. Validación y obtención de datos
        if isinstance(matricula_or_data, str):
            matricula = matricula_or_data
            estudiante_data = self.sistema.obtener_estudiante_por_matricula(matricula)
            if not estudiante_data:
                return f"❌ Error: No se encontraron datos para la matrícula {matricula}", None
        else:
            estudiante_data = matricula_or_data

        # 2. Buscar plantilla
        plantilla_path = os.path.join(self.base_path, "templates", "formato_calificaciones.docx")
        if not os.path.exists(plantilla_path):
            plantilla_path = os.path.join(self.base_path, "app", "formato_calificaciones.docx")

        try:
            doc = Document(plantilla_path)
        except FileNotFoundError:
            return f"❌ Error: No se encontró la plantilla en {plantilla_path}", None

        # 3. Procesar datos (encabezado + materias)
        promedio_general = estudiante_data.get("promedio_general", 0)
        estatus = self.sistema.obtener_estatus_por_promedio(promedio_general)

        # Encabezado + promedio + estatus
        placeholders_header = {
            "{{NOMBRE_INSTITUCION}}": "Instituto Tecnológico X",  # cámbialo si quieres
            "[[NOMBRE_COMPLETO]]": estudiante_data.get("nombre_completo", "N/A"),
            "{{NUMERO_MATRICULA}}": estudiante_data.get("matricula", "N/A"),
            "<<FECHA_EMISION>>": datetime.now().strftime("%d/%m/%Y"),
            "##PROMEDIO_GENERAL##": f"{promedio_general:.2f}",
            "##STATUS##": estatus,
        }

        reemplazar_en_documento(doc, placeholders_header)

        # Tabla de materias (plantilla con Matemáticas y Programación)
        if doc.tables:
            tabla = doc.tables[0]
            materias = estudiante_data.get("materias", [])

            def promedio_materia(m):
                return round(
                    (
                        m["calificacion_parcial1"]
                        + m["calificacion_parcial2"]
                        + m["calificacion_parcial3"]
                    ) / 3,
                    2,
                )

            # Matemáticas: primera materia
            if len(materias) > 0:
                mat = materias[0]
                placeholders_mat = {
                    "{{MAT_P1}}": str(mat["calificacion_parcial1"]),
                    "{{MAT_P2}}": str(mat["calificacion_parcial2"]),
                    "{{MAT_P3}}": str(mat["calificacion_parcial3"]),
                    "{{MAT_PROMEDIO}}": str(promedio_materia(mat)),
                }
                reemplazar_en_documento(doc, placeholders_mat)

            # Programación: segunda materia
            if len(materias) > 1:
                prog = materias[1]
                placeholders_prog = {
                    "{{PROG_P1}}": str(prog["calificacion_parcial1"]),
                    "{{PROG_P2}}": str(prog["calificacion_parcial2"]),
                    "{{PROG_P3}}": str(prog["calificacion_parcial3"]),
                    "{{PROG_PROMEDIO}}": str(promedio_materia(prog)),
                }
                reemplazar_en_documento(doc, placeholders_prog)

        # 4. Guardar archivo y RETORNAR RUTA
        nombre_archivo = f"reporte_{estudiante_data.get('matricula')}.docx"
        doc.save(nombre_archivo)

        ruta_absoluta = os.path.abspath(nombre_archivo)
        return f"✅ Formato generado: {nombre_archivo}", ruta_absoluta
