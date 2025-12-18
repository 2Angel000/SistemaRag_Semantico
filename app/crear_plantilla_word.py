from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def crear_plantilla_word(nombre_archivo="templates/formato_calificaciones.docx"):
    """Crea un documento Word con placeholders y una tabla para calificaciones."""
    document = Document()

    # Título
    document.add_heading("REPORTE DE CALIFICACIONES", level=0).alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    # Información general del estudiante
    document.add_paragraph("Fecha de Emisión: {{fecha_emision}}")
    document.add_paragraph("Nombre Completo: {{nombre_completo}}")
    document.add_paragraph("Matrícula: {{matricula}}")
    document.add_paragraph("Carrera: {{carrera}}")
    document.add_paragraph("Promedio General: {{promedio_general}}")
    document.add_paragraph("Estatus: {{estatus}}")
    document.add_paragraph("\n")  # Espacio en blanco

    # Encabezado de la tabla de materias
    document.add_heading("Detalle de Calificaciones por Materia", level=1)
    table = document.add_table(rows=1, cols=7)
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Materia"
    hdr_cells[1].text = "P1"
    hdr_cells[2].text = "P2"
    hdr_cells[3].text = "P3"
    hdr_cells[4].text = "Promedio"
    hdr_cells[5].text = "Asistencias"
    hdr_cells[6].text = "Faltas"

    document.save(nombre_archivo)
    print(f"Plantilla '{nombre_archivo}' creada exitosamente con 7 columnas.")


if __name__ == "__main__":
    crear_plantilla_word()
